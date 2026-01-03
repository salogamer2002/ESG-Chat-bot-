import os
from dotenv import load_dotenv
import fitz  # PyMuPDF
from sentence_transformers import SentenceTransformer
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from app.file_analysis import extract_text_by_page
from app.utils import chunk_text
from app.ingest import get_all_files
import json
import datetime

from tqdm import tqdm

from docx import Document as DocxDocument

load_dotenv()

RUN_LOG_DIR = "metadata_logs"
os.makedirs(RUN_LOG_DIR, exist_ok=True)

# Use timestamped file for each run
run_id = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
log_filename = f"{run_id}_metadata.jsonl"
log_path = os.path.join(RUN_LOG_DIR, log_filename)

# os.makedirs("rag_logs", exist_ok=True)
# run_id = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
# metadata_log_path = f"rag_logs/embedding_metadata_{run_id}.jsonl"

llm = ChatOpenAI(temperature=0, model="gpt-3.5-turbo")  # Or 3.5

# metadata_prompt = PromptTemplate.from_template("""
# You are a document classifier for an ESG compliance system.

# Below is a sample of a legal, regulatory, or ESG-related document.
# Identify the most likely metadata fields:

# 1. Jurisdiction (e.g., EU, Jordan, Kenya, China, US, Global)
# 2. Document Type (choose one of: law, regulation, report, policy, framework, standard, code, faq, publication, guidance)
# 3. Key Themes (comma-separated, e.g., labor rights, climate disclosure, governance, sustainable finance, human rights)

# Respond in valid JSON with fields: `jurisdiction`, `document_type`, `themes`.

# ---
# {document_excerpt}
# ---
# """)
from langchain.prompts import PromptTemplate

metadata_prompt = PromptTemplate.from_template("""
You are a document classifier for an ESG compliance system.

You are given:
- A sample excerpt from a legal, regulatory, or ESG-related document.
- Folder-level classification (helpful, but not definitive): `{folder_document_type}`

Your task is to extract the following metadata based on both content and context:

1. **Jurisdiction** – The most likely geographical or regulatory scope (e.g., EU, Jordan, Kenya, China, US, Global)
2. **Document Type** – The kind of document this is. Examples include: law, regulation, report, guideline, policy, framework, standard, code, FAQ, publication, etc. Use your judgment; you are not limited to these.
3. **Key Themes** – Comma-separated key focus areas of the document. Examples include: labor rights, climate disclosure, governance, sustainable finance, human rights, ESG integration, supply chain due diligence, biodiversity, gender equality.

Return your response as **valid JSON** with the fields:
- `jurisdiction`
- `document_type`
- `themes`

---
{document_excerpt}
---
""")

metadata_chain = LLMChain(llm=llm, prompt=metadata_prompt)

# Load previously saved metadata logs (across all files)
def load_previous_metadata(log_dir):
    previous = {}
    for fname in os.listdir(log_dir):
        if fname.endswith(".jsonl"):
            with open(os.path.join(log_dir, fname), "r", encoding="utf-8") as f:
                for line in f:
                    try:
                        item = json.loads(line)
                        if "source" in item and "llm_document_type" in item:
                            previous[item["source"]] = item
                    except:
                        continue
    return previous

cached_metadata_by_path = load_previous_metadata(RUN_LOG_DIR)
print(f"[CACHE LOAD] Found {len(cached_metadata_by_path)} cached metadata records.")

def clean_metadata(metadata: dict) -> dict:
    return {
        k: v for k, v in metadata.items()
        if v not in ("unspecified", "unknown", "", None, [], {})
    }

def extract_text_from_docx(path):
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_md(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()
    
def get_meaningful_excerpt(pages, max_chars=3000):
    filtered_pages = []
    for num, text in pages[:5]:
        if (
            len(text) > 200 and
            not ("table of contents" in text.lower() or "contents" in text.lower()) and
            text.count('.') / max(len(text), 1) < 0.05  # crude filter for dotted TOCs
        ):
            filtered_pages.append(text)
        if sum(len(p) for p in filtered_pages) >= max_chars:
            break
    return "\n".join(filtered_pages)[:max_chars]

def infer_metadata_from_path(filepath: str) -> dict:
    path_parts = filepath.lower().split(os.sep)

    metadata = {
        "jurisdiction": "unknown",
        "document_type": "unspecified",
        "framework": "",
        "tags": []
    }

    # Deepest folder → top-level
    for part in reversed(path_parts):
        part_clean = part.replace("_", " ").replace("-", " ").lower()

        # Specific document types first
        if "faq" in part_clean or "questions and answers" in part_clean:
            metadata["document_type"] = "faq"
            break
        elif "guiding document" in part_clean or "guidance" in part_clean:
            metadata["document_type"] = "guidance"
            break
        elif "standard" in part_clean:
            metadata["document_type"] = "standard"
            break
        elif "report" in part_clean:
            metadata["document_type"] = "report"
            break
        elif "code" in part_clean:
            metadata["document_type"] = "code"
            break
        elif "law" in part_clean or "legislation" in part_clean:
            metadata["document_type"] = "law"
            break
        elif "book" in part_clean or "paper" in part_clean:
            metadata["document_type"] = "publication"
            break

    # Broader metadata based on top → bottom
    for part in path_parts:
        if "frameworks" in part:
            if metadata["document_type"] in ("unspecified", ""):
                metadata["document_type"] = "framework"

            if "esrs" in part: metadata["framework"] = "ESRS"
            elif "gri" in part: metadata["framework"] = "GRI"
            elif "ifrs" in part: metadata["framework"] = "IFRS"
            elif "sasb" in part: metadata["framework"] = "SASB"
            elif "tcfd" in part: metadata["framework"] = "TCFD"
            elif "ghg" in part: metadata["framework"] = "GHG"

            # if "eu" in part:
            #     metadata["jurisdiction"] = "EU"

        elif "guiding reports" in part:
            metadata["document_type"] = "report"

        elif "legislation" in part or "law" in part:
            metadata["document_type"] = "law"

        elif "organization report" in part:
            metadata["document_type"] = "report"
            metadata["tags"].append("organization_report")

        elif "sectorassociation" in part:
            metadata["tags"].append("sector_code")

        # Jurisdictional hints
        # Normalize and tokenize all path parts
        # path_parts_lower = [p.lower() for p in path_parts]
        # all_parts_joined = " ".join(path_parts_lower)

        # if any("china" in p for p in path_parts_lower):
        #     metadata["jurisdiction"] = "China"
        # elif any("us" == p or "u.s." in p or "u.s" == p for p in path_parts_lower):
        #     if any("california" in p for p in path_parts_lower):
        #         metadata["jurisdiction"] = "US-California"
        #     elif any("minnesota" in p for p in path_parts_lower):
        #         metadata["jurisdiction"] = "US-Minnesota"
        #     else:
        #         metadata["jurisdiction"] = "US"
        # elif any("csrd" in p or "csddd" in p for p in path_parts_lower):
        #     metadata["jurisdiction"] = "EU"
        # elif "eu" in all_parts_joined:
        #     metadata["jurisdiction"] = "EU"

        # Tagging by special folder groups
        for tag in ["csrd", "sfdr", "csddd", "ilo", "sdgs", "ndrd", "issb"]:
            if tag in part:
                metadata["tags"].append(tag.upper())

    return metadata



def embed_documents(data_path="data/raw_docs", index_path="vector_store/faiss_index"):
    counter = 0
    print("Starting")
    print(datetime.datetime.now())
    model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    docs = []

    with open(log_path, "w", encoding="utf-8") as metadata_log:
        for filepath in get_all_files(data_path):
            if filepath.endswith(".pdf") and not os.path.basename(filepath).startswith("~$"):
                counter = counter + 1
                print(f"{counter}. Processing {filepath}")
                print(datetime.datetime.now())
                try:
                    with fitz.open(filepath) as pdf:
                        text = "\n".join([page.get_text() for page in pdf])
                except Exception as e:
                    print(f"[ERROR] Skipping PDF {filepath}: {e}")
                    print(datetime.datetime.now())
                    continue
            elif filepath.endswith(".docx") and not os.path.basename(filepath).startswith("~$"):
                counter = counter + 1
                print(f"{counter}. Processing {filepath}")
                print(datetime.datetime.now())
                try:
                    text = extract_text_from_docx(filepath)
                except Exception as e:
                    print(f"[ERROR] Skipping DOCX {filepath}: {e}")
                    print(datetime.datetime.now())
                    continue
            else:
                counter = counter + 1
                print("\n%s\n%s\n%s", "*" * 100, f" Skipping {filepath} ".center(100), "*" * 100)
                print(datetime.datetime.now())
                continue  # skip non-supported files

            folder_metadata = infer_metadata_from_path(filepath)
            # print("------------------- folder_metadata -------------------")
            # folder_metadata = clean_metadata(folder_metadata)
            # print(folder_metadata)

            if filepath in cached_metadata_by_path:
                metadata_fields = {
                    "jurisdiction": cached_metadata_by_path[filepath].get("llm_jurisdiction", "unknown"),
                    "document_type": cached_metadata_by_path[filepath].get("llm_document_type", "unspecified"),
                    "themes": cached_metadata_by_path[filepath].get("llm_themes", ""),
                }
                print("[CACHE HIT] Using previous LLM metadata.")
            else:
                pages = extract_text_by_page(filepath)
                intro = get_meaningful_excerpt(pages)
                # print("------------------- intro -------------------")
                # print(intro)
                
                response = metadata_chain.run(document_excerpt=intro, folder_document_type=folder_metadata.get("document_type", "unspecified"))
                # print("------------------- response -------------------")
                # print(response)

                try:
                    metadata_fields = json.loads(response)
                    # metadata_fields = clean_metadata(metadata_fields)
                except Exception:
                    print(f"[WARNING] LLM returned invalid JSON for {filepath}")
                    metadata_fields = {"jurisdiction": "unknown", "document_type": "unspecified", "themes": ""}
                # print("------------------- metadata_fields -------------------")
                # print(metadata_fields)

            chunks = chunk_text(text, max_tokens=512)
    # Clean and prefix LLM metadata
            llm_metadata = {f"llm_{k}": v for k, v in metadata_fields.items()}
            llm_metadata = clean_metadata(llm_metadata)

            # Clean and prefix folder metadata
            folder_metadata = {f"folder_{k}": v for k, v in folder_metadata.items()}
            folder_metadata = clean_metadata(folder_metadata)

            for i, chunk in enumerate(chunks):
                # Additional metadata from file/chunk
                chunk_meta = {
                    "file_name": os.path.basename(filepath),
                    "file_extension": os.path.splitext(filepath)[1][1:].lower(),
                    "chunk_id": f"{os.path.basename(filepath)}_chunk_{i}",
                    "word_count": len(chunk.split())
                }

                # Optional: add language detection if needed
                # from langdetect import detect
                # try:
                #     chunk_meta["language"] = detect(chunk)
                # except Exception:
                #     chunk_meta["language"] = "unknown"

                combined_metadata = {
                    "source": filepath,
                    "chunk": i,
                    **chunk_meta,
                    **llm_metadata,
                    **folder_metadata
                }

                # print("------------------- combined_metadata -------------------")
                # print(combined_metadata)


                docs.append(Document(page_content=chunk, metadata=combined_metadata))

                metadata_log.write(json.dumps(combined_metadata, ensure_ascii=False) + "\n")
                
            # if counter > 30:
            #     break

    print("Embeddings Generated")
    print(datetime.datetime.now())
    db = FAISS.from_documents(docs, model)
    db.save_local(index_path)
    print("Saved")
    print(datetime.datetime.now())
    print("Exit")


def add_single_document_to_faiss(filepath: str, index_path="vector_store/faiss_index"):
    # Load FAISS index
    model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.load_local(index_path, model, allow_dangerous_deserialization=True)

    # Extract text
    try:
        if filepath.endswith(".pdf"):
            with fitz.open(filepath) as pdf:
                text = "\n".join([page.get_text() for page in pdf])
        elif filepath.endswith(".docx"):
            text = extract_text_from_docx(filepath)
        elif filepath.endswith(".md"):
            text = extract_text_from_md(filepath)
        else:
            raise ValueError("Unsupported file type")
    except Exception as e:
        print(f"[ERROR] Failed to extract text: {e}")
        return

    # Infer metadata
    folder_metadata = infer_metadata_from_path(filepath)
    
    if filepath.endswith(".md"):
        pages = [(0, text)]
    else:
        pages = extract_text_by_page(filepath)

    intro = get_meaningful_excerpt(pages)
    try:
        response = metadata_chain.run(
            document_excerpt=intro,
            folder_document_type=folder_metadata.get("document_type", "unspecified")
        )
        metadata_fields = json.loads(response)
    except Exception as e:
        print(f"[WARNING] LLM metadata failed: {e}")
        metadata_fields = {
            "jurisdiction": "unknown",
            "document_type": "unspecified",
            "themes": ""
        }

    # Clean metadata
    llm_metadata = {f"llm_{k}": v for k, v in metadata_fields.items()}
    llm_metadata = clean_metadata(llm_metadata)
    folder_metadata = {f"folder_{k}": v for k, v in folder_metadata.items()}
    folder_metadata = clean_metadata(folder_metadata)

    # Chunk and embed
    chunks = chunk_text(text, max_tokens=512)
    docs = []
    for i, chunk in enumerate(chunks):
        chunk_meta = {
            "file_name": os.path.basename(filepath),
            "file_extension": os.path.splitext(filepath)[1][1:].lower(),
            "chunk_id": f"{os.path.basename(filepath)}_chunk_{i}",
            "word_count": len(chunk.split())
        }

        combined_metadata = {
            "source": filepath,
            "chunk": i,
            **chunk_meta,
            **llm_metadata,
            **folder_metadata
        }

        docs.append(Document(page_content=chunk, metadata=combined_metadata))

    # Add to vector store and save
    db.add_documents(docs)
    db.save_local(index_path)
    print(f"[DONE] Successfully added {len(docs)} chunks from {filepath} to FAISS index.")

    # Optional: log metadata
    log_dir = "metadata_logs"
    os.makedirs(log_dir, exist_ok=True)
    log_file = os.path.join(log_dir, f"{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_single_add.jsonl")
    with open(log_file, "w", encoding="utf-8") as f:
        for doc in docs:
            f.write(json.dumps(doc.metadata, ensure_ascii=False) + "\n")
